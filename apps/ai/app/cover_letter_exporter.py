from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from docx import Document


def generate_cover_letter_pdf_bytes(content: str) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4
    left_margin = 20 * mm
    top_margin = height - 20 * mm
    line_height = 6 * mm

    text_object = pdf.beginText()
    text_object.setTextOrigin(left_margin, top_margin)
    text_object.setLeading(line_height)
    text_object.setFont("Times-Roman", 11)

    current_y = top_margin
    for raw_line in content.splitlines():
        line = raw_line.rstrip()

        if not line:
            current_y -= line_height
            text_object.textLine("")
            continue

        # Basic wrap
        words = line.split()
        current_line = ""

        for word in words:
            test_line = f"{current_line} {word}".strip()
            if pdf.stringWidth(test_line, "Times-Roman", 11) <= (width - 40 * mm):
                current_line = test_line
            else:
                text_object.textLine(current_line)
                current_y -= line_height
                current_line = word

                if current_y <= 20 * mm:
                    pdf.drawText(text_object)
                    pdf.showPage()
                    text_object = pdf.beginText()
                    text_object.setTextOrigin(left_margin, top_margin)
                    text_object.setLeading(line_height)
                    text_object.setFont("Times-Roman", 11)
                    current_y = top_margin

        if current_line:
            text_object.textLine(current_line)
            current_y -= line_height

        if current_y <= 20 * mm:
            pdf.drawText(text_object)
            pdf.showPage()
            text_object = pdf.beginText()
            text_object.setTextOrigin(left_margin, top_margin)
            text_object.setLeading(line_height)
            text_object.setFont("Times-Roman", 11)
            current_y = top_margin

    pdf.drawText(text_object)
    pdf.save()
    buffer.seek(0)
    return buffer.read()


def generate_cover_letter_docx_bytes(content: str) -> bytes:
    doc = Document()

    for line in content.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()